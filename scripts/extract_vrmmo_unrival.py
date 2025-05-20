import pandas as pd
from docx import Document
from docx.shared import Pt
import sys


def process_excel_to_word(input_file, output_file):
    # Load the Excel file
    sheet_data = pd.read_excel(input_file, sheet_name='Sheet1')

    # Process the data: group by "Header" but keep individual rows as separate paragraphs
    processed_data = sheet_data.groupby('Header')['Body'].apply(list).reset_index()

    # Sort chapters by the header text for logical order
    processed_data = processed_data.sort_values(by='Header')

    # Create a Word document
    doc = Document()

    # Add content to the Word document
    for _, row in processed_data.iterrows():
        # Add the header with specified styling
        header = doc.add_heading(row['Header'], level=1)
        header.style.font.name = 'Arial'
        header.style.font.size = Pt(20)

        # Add each row as a new paragraph
        for paragraph in row['Body']:
            if isinstance(paragraph, str) and paragraph.strip():  # Ensure it's a valid string
                body = doc.add_paragraph(paragraph.strip())
                body.style.font.name = 'Arial'
                body.style.font.size = Pt(11)

        # Add a page break after each chapter
        doc.add_page_break()

    # Save the document
    doc.save(output_file)
    print(f"Document saved at: {output_file}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python script_name.py <path_to_input_excel_file> <path_to_output_word_file>")
    else:
        input_file = sys.argv[1]
        output_file = sys.argv[2]
        process_excel_to_word(input_file, output_file)
